"""
Document Formatter

Converts documents to various formats.
"""

from typing import Optional
from pathlib import Path
from .document import Document, DocumentFormat


class DocumentFormatter:
    """Formats documents to various file formats."""
    
    def __init__(self, document: Document):
        """
        Initialize document formatter.
        
        Args:
            document: Document to format
        """
        self.document = document
    
    async def to_docx(self, output_path: str) -> str:
        """
        Export document to DOCX format.
        
        Args:
            output_path: Output file path
        
        Returns:
            Path to created file
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = DocxDocument()
        
        # Add title
        doc.add_heading(self.document.title, 0)
        
        # Add paragraphs
        for para in self.document.paragraphs:
            p = doc.add_paragraph(para.text)
            
            # Apply formatting
            if para.bold:
                p.runs[0].bold = True
            if para.italic:
                p.runs[0].italic = True
            if para.underline:
                p.runs[0].underline = True
            
            if para.font_size:
                p.runs[0].font.size = Pt(para.font_size)
            
            if para.alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif para.alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif para.alignment == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add tables
        for table in self.document.tables:
            doc_table = doc.add_table(rows=len(table.rows), cols=len(table.rows[0]) if table.rows else 0)
            
            # Add headers
            if table.headers:
                header_cells = doc_table.rows[0].cells
                for i, header in enumerate(table.headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Add rows
            start_row = 1 if table.headers else 0
            for i, row in enumerate(table.rows):
                row_cells = doc_table.rows[start_row + i].cells
                for j, cell in enumerate(row):
                    row_cells[j].text = str(cell)
        
        # Add images
        for img in self.document.images:
            # Download and add image (simplified)
            # In production, handle image download properly
            pass
        
        doc.save(output_path)
        return output_path
    
    async def to_pdf(self, output_path: str) -> str:
        """
        Export document to PDF format.
        
        Args:
            output_path: Output file path
        
        Returns:
            Path to created file
        """
        # First convert to DOCX, then to PDF
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            await self.to_docx(tmp_path)
            
            # Convert DOCX to PDF using LibreOffice or similar
            # This is simplified - in production use proper PDF conversion
            try:
                import subprocess
                subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(Path(output_path).parent),
                    tmp_path
                ], check=True)
                
                pdf_path = Path(tmp_path).with_suffix(".pdf")
                if pdf_path.exists():
                    import shutil
                    shutil.move(str(pdf_path), output_path)
            except Exception:
                # Fallback: return DOCX path if PDF conversion fails
                import shutil
                shutil.copy(tmp_path, output_path.replace(".pdf", ".docx"))
                raise Exception("PDF conversion failed. LibreOffice required.")
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        
        return output_path
    
    async def to_html(self, output_path: str) -> str:
        """
        Export document to HTML format.
        
        Args:
            output_path: Output file path
        
        Returns:
            Path to created file
        """
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{self.document.title}</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "h1 { color: #333; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f2f2f2; }",
            "img { max-width: 100%; height: auto; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{self.document.title}</h1>"
        ]
        
        # Add paragraphs
        for para in self.document.paragraphs:
            style_parts = []
            if para.bold:
                style_parts.append("font-weight: bold")
            if para.italic:
                style_parts.append("font-style: italic")
            if para.underline:
                style_parts.append("text-decoration: underline")
            if para.font_size:
                style_parts.append(f"font-size: {para.font_size}pt")
            if para.color:
                style_parts.append(f"color: {para.color}")
            if para.alignment:
                style_parts.append(f"text-align: {para.alignment}")
            
            style_attr = f' style="{"; ".join(style_parts)}"' if style_parts else ""
            html_parts.append(f"<p{style_attr}>{para.text}</p>")
        
        # Add tables
        for table in self.document.tables:
            html_parts.append("<table>")
            if table.headers:
                html_parts.append("<thead><tr>")
                for header in table.headers:
                    html_parts.append(f"<th>{header}</th>")
                html_parts.append("</tr></thead>")
            html_parts.append("<tbody>")
            for row in table.rows:
                html_parts.append("<tr>")
                for cell in row:
                    html_parts.append(f"<td>{cell}</td>")
                html_parts.append("</tr>")
            html_parts.append("</tbody></table>")
        
        # Add images
        for img in self.document.images:
            html_parts.append(f'<img src="{img.url}" alt="{img.caption or ""}" />')
            if img.caption:
                html_parts.append(f"<p><em>{img.caption}</em></p>")
        
        html_parts.extend(["</body>", "</html>"])
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_parts))
        
        return output_path
    
    async def to_markdown(self, output_path: str) -> str:
        """
        Export document to Markdown format.
        
        Args:
            output_path: Output file path
        
        Returns:
            Path to created file
        """
        md_parts = [f"# {self.document.title}\n"]
        
        # Add paragraphs
        for para in self.document.paragraphs:
            text = para.text
            if para.bold:
                text = f"**{text}**"
            if para.italic:
                text = f"*{text}*"
            md_parts.append(text)
            md_parts.append("")  # Empty line
        
        # Add tables
        for table in self.document.tables:
            if table.headers:
                md_parts.append("| " + " | ".join(table.headers) + " |")
                md_parts.append("| " + " | ".join(["---"] * len(table.headers)) + " |")
            for row in table.rows:
                md_parts.append("| " + " | ".join(str(cell) for cell in row) + " |")
            md_parts.append("")
        
        # Add images
        for img in self.document.images:
            md_parts.append(f"![{img.caption or ''}]({img.url})")
            md_parts.append("")
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_parts))
        
        return output_path

