"""
Document Indexer

Indexes various document types (PDF, DOCX, TXT, Markdown).
"""

import os
from typing import List, Optional
from datetime import datetime
from pathlib import Path
import mimetypes

from .models import IndexedDocument, IndexMetadata, DocumentType


class DocumentIndexer:
    """Indexes documents from files."""
    
    def __init__(self):
        """Initialize document indexer."""
        pass
    
    async def index_file(self, file_path: str) -> IndexedDocument:
        """
        Index a document file.
        
        Args:
            file_path: Path to the file
        
        Returns:
            IndexedDocument
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine document type
        doc_type = self._detect_type(file_path)
        
        # Extract text based on type
        content = await self._extract_text(file_path, doc_type)
        
        # Create chunks
        chunks = self._chunk_text(content)
        
        # Create metadata
        metadata = IndexMetadata(
            document_id=f"doc_{path.stem}_{path.stat().st_mtime}",
            document_type=doc_type,
            source=str(path.absolute()),
            title=path.stem,
            size=path.stat().st_size,
            indexed_at=datetime.utcnow()
        )
        
        return IndexedDocument(
            id=metadata.document_id,
            content=content,
            metadata=metadata,
            chunks=chunks
        )
    
    def _detect_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension."""
        ext = Path(file_path).suffix.lower()
        
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    async def _extract_text(self, file_path: str, doc_type: DocumentType) -> str:
        """Extract text from document based on type."""
        if doc_type == DocumentType.TXT or doc_type == DocumentType.MARKDOWN:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif doc_type == DocumentType.PDF:
            # Use PyPDF2 or pdfplumber
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                raise ImportError("PyPDF2 is required for PDF extraction. Install with: pip install PyPDF2")
        elif doc_type == DocumentType.DOCX:
            # Use python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into chunks."""
        chunks = []
        words = text.split()
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            if current_length + word_length > chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                # Start new chunk with overlap
                overlap_words = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

